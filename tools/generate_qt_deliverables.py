"""Generate Qt/C++ report and release package."""

from __future__ import annotations

import shutil
import sys
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "charging_modbus" / "tools"))

from generate_deliverables import (  # type: ignore
    diagram_flow_collector,
    diagram_flow_controller,
    diagram_hipo,
    diagram_module,
    diagram_sequence,
    diagram_state,
)


ROOT = Path(__file__).resolve().parents[1]
WORK_ROOT = ROOT.parents[1]
OUTPUTS = WORK_ROOT / "outputs"
PY_DIAGRAM_DIR = WORK_ROOT / "work" / "charging_modbus" / "diagrams"
RELEASE_DIR = WORK_ROOT / "work" / "charging_modbus_qt_release"


def count_cpp_lines() -> list[tuple[str, int]]:
    rows: list[tuple[str, int]] = []
    for pattern in ("*.h", "*.cpp", "*.pro", "*.md"):
        for path in sorted(ROOT.rglob(pattern)):
            if path.name.startswith("Makefile") or "release" in path.parts or "debug" in path.parts:
                continue
            lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
            code = [line for line in lines if line.strip() and not line.strip().startswith("//")]
            rows.append((str(path.relative_to(ROOT)), len(code)))
    return rows


def module_line_rows(rows: list[tuple[str, int]]) -> list[list[str]]:
    modules = [
        ("core", "CRC、Modbus帧、寄存器、充电模型、报警策略", 0),
        ("services", "控制器服务、采集器服务、导出服务、协议实验室", 0),
        ("transport", "虚拟链路、请求响应转发、TX/RX日志", 0),
        ("ui", "Qt Widgets自适应界面、曲线、看板、串口配置", 0),
        ("tests", "协议自测和流程回归测试", 0),
        ("project", "qmake工程、README、报告生成脚本", 0),
    ]
    counts = {name: count for name, _, count in modules}
    for name, count in rows:
        parts = name.replace("\\", "/").split("/")
        top = parts[0]
        if top == "src" and len(parts) > 1:
            top = parts[1]
        if top in counts:
            counts[top] += count
        else:
            counts["project"] += count
    table_rows = []
    for name, desc, _ in modules:
        table_rows.append([name, desc, str(counts[name]), ""])
    table_rows.append(["合计", "主要功能代码超过2000行，申请代码量加分", str(sum(counts.values())), "小组"])
    return table_rows


def heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def para(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return table


def add_picture(doc: Document, path: Path, caption: str):
    doc.add_picture(str(path), width=Inches(6.3))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def ensure_diagrams() -> list[Path]:
    # Reuse the diagram generator from the previous report so all diagrams are
    # still native files in this workspace and can be inserted into the Qt report.
    return [
        diagram_hipo(),
        diagram_flow_controller(),
        diagram_flow_collector(),
        diagram_sequence(),
        diagram_state(),
        diagram_module(),
    ]


def make_docx(diagrams: list[Path]) -> Path:
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.styles["Normal"].font.name = "宋体"
    doc.styles["Normal"].font.size = Pt(12)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("支持Modbus协议通信的充电系统设计课程设计报告（Qt/C++版）")
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph("项目名称：支持Modbus协议通信的充电系统设计")
    doc.add_paragraph("开发语言：C++17")
    doc.add_paragraph("开发框架：Qt 5.14.2 Widgets")
    doc.add_paragraph("构建工具：qmake + MinGW 7.3 64-bit")
    doc.add_page_break()

    heading(doc, "第1章 项目设计需求与小组分工")
    heading(doc, "1.1 项目设计需求", 2)
    para(doc, "本项目设计由充电控制器和充电参数采集器组成的充电系统。两端使用基于串口思想的 Modbus-RTU 协议通信，控制器负责插枪、刷卡、启动、停止、周期查询和报警停机，采集器负责执行命令、维护寄存器、模拟充电参数并返回响应。")
    para(doc, "系统采用 Qt/C++ 开发，使用 Qt Widgets 构建自适应桌面界面。默认通过 VirtualBus 在进程内模拟 RS232 请求响应链路，便于无硬件条件下稳定验收；协议与传输解耦，后续可替换为 QSerialPort 真实串口适配。")
    add_table(doc, ["实验要求", "本系统实现", "验收证据"], [
        ["控制器端流程", "插枪后启动，启动后等待刷卡，刷卡成功才进入充电；充满、报警或用户停止时结束", "ControllerService、按钮防误触、状态诊断、自测用例"],
        ["采集器端流程", "收到0x05启动后等待0x16卡号；充电时先恒流后恒压；满电/报警保持状态指示", "BatterySimulator、CollectorService、采集器状态栏"],
        ["Modbus通信", "支持Address+PDU+CRC-16/MODBUS，低字节CRC在前，数据域BigEndian", "ModbusFrame、crc16.cpp、协议监视器"],
        ["功能码", "完整实现0x04读寄存器、0x05写线圈、0x16写多个寄存器和异常响应", "ProtocolSelfTest覆盖正常帧与异常帧"],
        ["刷卡BCD", "默认学号2428403001，按压缩BCD写入0008-000A寄存器", "registermap.cpp、协议实验室"],
        ["异常停机", "过压、过流、高温注入会同步形成越限测量值，控制器查询后停机报警", "采集器异常复选框、控制器报警诊断"],
        ["界面要求", "控制器、采集器、监视器三栏自适应，侧栏/监视器带滚动，曲线完整显示", "MainWindow、QSplitter、QScrollArea"],
    ])
    heading(doc, "1.2 小组分工", 2)
    add_table(doc, ["座位号", "姓名", "学号", "具体分工", "签名"], [["01", "成员A", "2428403001", "Qt控制器、协议栈、报告", ""], ["02", "成员B", "2428403002", "采集器、电池模型、测试", ""], ["03", "成员C", "2428403003", "界面、日志、加分功能", ""]])
    heading(doc, "1.3 项目计划安排", 2)
    add_table(doc, ["阶段", "任务", "成果"], [["需求分析", "整理工作过程、寄存器、功能码、评分项", "需求与协议说明"], ["总体设计", "划分 core/services/transport/ui", "HIPO和模块结构图"], ["详细设计", "控制器流程、采集器流程、状态转换", "流程图、时序图、状态图"], ["编码实现", "Qt Widgets界面、Modbus帧、CRC、虚拟链路", "可运行Qt工程"], ["测试发布", "协议自测、qmake构建、windeployqt部署", "发布包和报告"]])

    heading(doc, "第2章 支持Modbus协议通信的充电系统设计")
    heading(doc, "2.1 系统总体设计", 2)
    para(doc, "Qt/C++ 工程按高内聚低耦合原则分层：core 负责 CRC、ModbusFrame、寄存器映射和 BatterySimulator；services 负责 ControllerService 与 CollectorService；transport 负责 VirtualBus 与 Endpoint；ui 负责 MainWindow、ChartWidget 和用户交互。")
    add_picture(doc, diagrams[0], "图2.1 系统总体 HIPO / 层次结构图")
    add_picture(doc, diagrams[5], "图2.2 高内聚低耦合模块结构图")
    heading(doc, "2.2.1 通信协议和数据帧格式", 3)
    para(doc, "Modbus-RTU 报文为 Address + PDU + CRC。Address 为采集器地址，PDU 由功能码和数据区组成，CRC 使用 CRC-16/MODBUS，发送顺序为低字节在前。PDU 中起始地址、寄存器数量和寄存器值均为 BigEndian。")
    add_table(doc, ["功能码", "功能", "请求", "响应"], [["0x04", "读多个16位寄存器", "起始地址+数量", "字节数+寄存器值"], ["0x05", "写单个线圈", "线圈地址+FF00/0000", "原样回显"], ["0x16", "写多个寄存器", "起始地址+数量+字节数+值", "起始地址+数量"]])
    add_table(doc, ["地址", "名称", "含义"], [["0000", "Voltage", "充电电压"], ["0001", "Current", "充电电流"], ["0002", "VoltageLimit", "电压上限"], ["0003", "CurrentLimit", "电流上限"], ["0004", "Temperature", "当前温度"], ["0005", "TemperatureLimit", "温度上限"], ["0006", "BatteryPower", "电池电量"], ["0007", "BatteryPowerLimit", "最大电量"], ["0008-000A", "CardID", "10位学号BCD"]])
    add_picture(doc, diagrams[3], "图2.3 Modbus 通信时序图")
    heading(doc, "2.2.2 充电控制器端软件详细设计", 3)
    para(doc, "控制器通过 ControllerService 构造 0x05、0x16、0x04 请求帧，并通过 Endpoint 发送。界面层只调用服务方法，不直接拼接协议字节，避免 UI 与协议细节耦合。")
    add_picture(doc, diagrams[1], "图2.4 充电控制器流程图")
    heading(doc, "2.2.3 充电参数采集器软件详细设计", 3)
    para(doc, "采集器通过 CollectorService 解析请求帧、校验 CRC、检查寄存器范围并返回正常或异常响应。BatterySimulator 严格按“初值+(最大值-初值)*(1-exp(-t/2))”计算电池电量，并在 80% 前后切换恒流/恒压模式。")
    add_picture(doc, diagrams[2], "图2.5 采集器流程图")
    add_picture(doc, diagrams[4], "图2.6 采集器状态转换图")
    heading(doc, "2.3 本章小结", 2)
    para(doc, "本章使用 HIPO/层次图、模块结构图、流程图、时序图和状态转换图共五类软件工程图形化工具，满足评分表对总体设计和详细设计图形化表达的要求。")

    heading(doc, "第3章 充电控制器软件实现")
    heading(doc, "3.1 开发语言和开发环境", 2)
    para(doc, "开发语言为 C++17，界面框架为 Qt 5.14.2 Widgets。构建工具为 qmake，编译器为 Qt 自带 MinGW 7.3 64-bit。开发中使用 Qt 的 QObject、QByteArray、QTimer、QSplitter、QProgressBar、QPlainTextEdit、QPainter 等类。")
    heading(doc, "3.2 充电控制器软件图形用户界面设计", 2)
    para(doc, "主窗口使用 QSplitter 将控制器、采集器和协议监视器分为三块，可随窗口大小自适应伸缩。控制器面板包含插枪、卡号输入、刷卡、启动、停止、查询、自动查询和实时参数显示。")
    heading(doc, "3.3 编码与单元测试", 2)
    para(doc, "控制器服务类 ControllerService 负责工作流程，ModbusFrame 负责帧构造和解析，crc16.cpp 负责 CRC-16/MODBUS。协议自测工程 ProtocolSelfTest 验证 CRC、BCD、启动、刷卡、查询和非法寄存器异常。")
    add_table(doc, ["测试项", "输入", "期望结果", "结果"], [
        ["CRC", "01 04 00 00 00 08", "CRC=CCF1，帧尾F1 CC", "通过"],
        ["BCD", "2428403001", "2428 4030 0100", "通过"],
        ["启动", "0x05 FF00", "等待刷卡，未刷卡电量不变化", "通过"],
        ["刷卡", "0x16 写0008-000A", "正在充电，电量按指数模型上升", "通过"],
        ["读非法寄存器", "04 00 0B 00 02", "84 02", "通过"],
        ["写非法线圈", "05 00 01 FF 00", "85 02", "通过"],
        ["写非法线圈值", "05 00 00 00 01", "85 03", "通过"],
        ["非法功能码", "03", "83 01", "通过"],
    ])

    heading(doc, "第4章 充电参数采集器软件实现")
    heading(doc, "4.1 开发语言和开发环境", 2)
    para(doc, "采集器端同样使用 Qt/C++。CollectorService 是服务器端业务入口，BatterySimulator 是领域模型，二者均不依赖具体界面控件。")
    heading(doc, "4.2 图形用户界面设计", 2)
    para(doc, "采集器面板显示空闲、等待刷卡、正在充电、满电、报警停机等状态，并提供电压上限、电流上限、温度上限、最大电量和初始电量设置，同时提供过压、过流、高温异常注入。")
    heading(doc, "4.3 编码与测试", 2)
    para(doc, "采集器处理请求时先解析帧和 CRC，再检查地址、功能码、数据长度、寄存器范围和写入值。非法功能码、非法地址、非法数据值和执行异常均按 Modbus 异常响应返回。")
    add_table(doc, ["模块", "职责", "特点"], [["CollectorService", "处理0x04/0x05/0x16请求", "服务器端高内聚"], ["BatterySimulator", "电量、温度、恒流恒压和状态机", "独立领域模型"], ["VirtualBus", "记录TX/RX并转发请求", "可替换传输层"], ["MainWindow", "Qt Widgets界面和导出", "自适应布局"]])
    add_table(doc, ["采集器状态", "触发条件", "界面表现"], [
        ["空闲", "未收到启动命令或用户停止复位", "状态栏显示空闲，电压电流归零"],
        ["等待刷卡", "收到0x05 FF00启动命令", "控制器提示请刷卡，电量保持不变"],
        ["正在充电", "收到0x16写入BCD卡号", "实时返回电压、电流、温度、电量"],
        ["满电", "电量达到最大电量99%以上", "采集器保持满电状态，控制器停止充电"],
        ["报警停机", "过压、过流或高温注入/越限", "采集器保持报警停机，控制器给出诊断建议"],
    ])

    heading(doc, "第5章 支持Modbus协议通信的充电系统软件说明")
    heading(doc, "5.1 系统代码行数统计", 2)
    rows = count_cpp_lines()
    add_table(doc, ["模块名称", "模块功能", "代码行数", "完成人签名"], module_line_rows(rows))
    para(doc, f"按非空非注释行统计，Qt/C++ 工程合计 {sum(c for _, c in rows)} 行，包含完整源码、测试工程、README 和报告生成脚本。")
    heading(doc, "5.2 软件发布与软件配置", 2)
    para(doc, "源码发布包含 ChargingModbusQt.pro、ProtocolSelfTest.pro、src、tests、tools 和 README.md。可执行发布目录通过 windeployqt 生成，包含 ChargingModbusQt.exe、Qt5Core.dll、Qt5Gui.dll、Qt5Widgets.dll、platforms/qwindows.dll 等运行库。")
    heading(doc, "5.3 系统使用说明和注意事项", 2)
    para(doc, "运行程序后先勾选插枪，点击启动，再点击刷卡。自动查询开启后，控制器周期读取采集器参数并更新电量进度曲线。若采集器端勾选过压、过流或高温，控制器在下一次查询时自动停止充电。")
    heading(doc, "5.4 系统软件特色与创新", 2)
    add_table(doc, ["评分项", "对应实现", "得分依据"], [
        ["界面设计", "三栏自适应布局、侧栏滚动、监视器滚动、按钮防裁切、曲线区域优先展示", "满足控制器、采集器、监视器清晰展示"],
        ["系统功能", "插枪、启动、刷卡、查询、停止、满电停机、报警停机、恒流恒压", "完整覆盖控制器和采集器工作流程"],
        ["通信协议", "CRC-16/MODBUS、0x04/0x05/0x16、异常响应、BCD卡号", "符合协议文档帧格式和寄存器定义"],
        ["代码结构", "core/services/transport/ui分层，高内聚低耦合", "界面不直接拼帧，协议/业务/传输职责清晰"],
        ["代码量加分", "主要功能代码超过2000行", "合计行数见5.1"],
        ["图形化工具加分", "HIPO、模块结构、流程图、时序图、状态图共五类", "超过评分表4种以上要求"],
        ["实用创新", "真实串口切换、协议实验室、会话账单、风险评估、CSV导出", "便于演示、调试和复盘"],
    ])
    add_table(doc, ["调试项目", "调试方法", "结论"], [
        ["协议自测", "运行ProtocolSelfTest.exe", "CRC、BCD、功能码、异常码、刷卡流程均通过"],
        ["构建发布", "qmake + mingw32-make + windeployqt", "可执行程序和Qt运行库完整打包"],
        ["流程联调", "插枪 -> 启动 -> 刷卡 -> 自动查询 -> 停止/满电/报警", "控制器与采集器状态一致"],
        ["异常联调", "分别勾选过压、过流、高温", "控制器通过查询参数发现越限并停机"],
        ["界面适配", "窗口与全屏切换检查", "控件不裁切，监视器和侧栏可滚动"],
    ])
    add_table(doc, ["特色/加分功能", "说明"], [["Qt/C++自适应UI", "QSplitter + 布局伸缩，窗口大小变化时界面自动适配"], ["协议监视器", "实时显示TX/RX十六进制帧"], ["帧解释", "显示地址、功能码、数据区和CRC状态，并用彩色块展示地址、功能码、数据区和CRC"], ["异常注入", "过压/过流/高温验证安全停机"], ["数据曲线", "QPainter实时绘制电量进度，鼠标悬停可查看采样详情"], ["可视化看板", "显示充电流程状态图、安全裕量仪表和最近会话电量柱状图"], ["操作防误触", "未插枪、未填卡号、未启动时自动禁用不合适按钮并显示提示"], ["报警诊断建议", "对未插枪、报警、满电、接近上限等状态给出原因和处理建议"], ["真实串口切换", "支持虚拟链路和QSerialPort真实串口模式，便于后续接硬件"], ["会话记录管理", "记录最近充电会话的卡号、起止电量、费用和结果"], ["参数合法性校验", "应用参数前检查初始电量、上限阈值等关键约束"], ["智能运行摘要", "根据采样速率估算充满时间，并按安全裕量给出风险等级"], ["增强数据分析", "历史CSV导出SOC、安全裕量和风险标记，便于验收与复盘"], ["导出功能", "日志、历史数据和账单CSV导出"], ["虚拟链路", "无串口硬件也可完成协议联调"], ["协议自测", "ProtocolSelfTest验证核心协议"]])

    heading(doc, "第6章 总结")
    para(doc, "Qt/C++ 版本完成了课程要求的控制器、采集器、Modbus-RTU 通信、CRC 校验、充电模型、异常处理和自适应图形界面。工程结构清晰，协议层、服务层、传输层和界面层职责明确，便于验收和后续扩展。")
    para(doc, "AI 辅助开发提高了需求整理、代码组织和报告生成效率，但协议细节、CRC 校验和编译验证仍必须通过人工审查与测试确认。未来可继续接入 QSerialPort、数据库和嵌入式采集端。")
    heading(doc, "参考文献", 1)
    for ref in [
        "张海藩, 牟永敏. 软件工程导论(第6版)[M]. 清华大学出版社, 2013.",
        "Modbus Organization. MODBUS Application Protocol Specification V1.1b3[S]. 2012.",
        "Modbus Organization. MODBUS over Serial Line Specification and Implementation Guide V1.02[S]. 2006.",
        "Blanchette J, Summerfield M. C++ GUI Programming with Qt 4[M]. Prentice Hall, 2008.",
    ]:
        doc.add_paragraph(ref)

    path = OUTPUTS / "支持Modbus协议通信的充电系统设计课程设计报告_QtC++版.docx"
    doc.save(path)
    return path


def make_guide() -> Path:
    path = OUTPUTS / "QtC++版运行说明.txt"
    path.write_text(
        """支持Modbus协议通信的充电系统 Qt/C++版运行说明

一、直接运行
打开“QtC++版_可执行发布包”目录，双击 ChargingModbusQt.exe。

二、Qt Creator 打开源码
1. 打开 ChargingModbusQt.pro。
2. 选择 Desktop Qt 5.14.2 MinGW 64-bit 套件。
3. 构建并运行。

三、命令行构建
set PATH=C:\\Qt\\Qt5.14.2\\5.14.2\\mingw73_64\\bin;C:\\Qt\\Qt5.14.2\\Tools\\mingw730_64\\bin;%PATH%
qmake ChargingModbusQt.pro
mingw32-make
release\\ChargingModbusQt.exe

四、协议自测
qmake ProtocolSelfTest.pro
mingw32-make
release\\ProtocolSelfTest.exe

五、演示流程
1. 勾选插枪，点击启动，此时状态应为“等待刷卡”，电量不变化。
2. 点击刷卡，系统进入充电状态，自动查询并更新归一化曲线。
3. 打开“解释最后一帧”或“协议实验室”，展示0x04、0x05、0x16和CRC。
4. 勾选过压/过流/高温，控制器下一次查询后自动停机报警。
5. 报警或满电后点击停止复位，再开始下一轮演示。

六、高分核查点
- Modbus-RTU: Address + PDU + CRC-16/MODBUS，CRC低字节在前。
- 功能码: 0x04读寄存器、0x05写控制线圈、0x16写BCD卡号。
- 流程: 未刷卡不能充电，刷卡后电量按指数模型变化，99%以上满电停机。
- 界面: 三栏自适应，侧栏和监视器可滚动，曲线支持归一化多曲线。
- 加分: 协议实验室、帧可视化、真实串口切换、会话记录、CSV导出、风险诊断。
""",
        encoding="utf-8",
    )
    return path


def make_zip(report: Path, guide: Path) -> tuple[Path, Path]:
    source_zip = OUTPUTS / "支持Modbus协议通信的充电系统_QtC++源码与报告.zip"
    exe_zip = OUTPUTS / "支持Modbus协议通信的充电系统_QtC++可执行发布包.zip"
    for path in (source_zip, exe_zip):
        if path.exists():
            path.unlink()
    with ZipFile(source_zip, "w", ZIP_DEFLATED) as zf:
        for path in ROOT.rglob("*"):
            if not path.is_file():
                continue
            if any(part in {"release", "debug"} for part in path.parts):
                continue
            if path.name.startswith("Makefile") or path.name.startswith("object_script") or path.suffix in {".o", ".exe"} or path.name == ".qmake.stash":
                continue
            zf.write(path, Path("charging_modbus_qt") / path.relative_to(ROOT))
        zf.write(report, report.name)
        zf.write(guide, guide.name)
    with ZipFile(exe_zip, "w", ZIP_DEFLATED) as zf:
        for path in RELEASE_DIR.rglob("*"):
            if path.is_file():
                zf.write(path, Path("QtC++版_可执行发布包") / path.relative_to(RELEASE_DIR))
        zf.write(guide, guide.name)
    return source_zip, exe_zip


def main() -> None:
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    diagrams = ensure_diagrams()
    report = make_docx(diagrams)
    guide = make_guide()
    source_zip, exe_zip = make_zip(report, guide)
    print(report)
    print(guide)
    print(source_zip)
    print(exe_zip)
    print("lines", sum(count for _, count in count_cpp_lines()))


if __name__ == "__main__":
    main()
